#!/usr/bin/env python3
"""Generate an editable .docx resume from a RenderCV cv.yaml, mirroring the
harvard theme: serif, centered ruled section titles, justified body, right-aligned
dates, tight consistent spacing.

Font is Georgia (not XCharter/Charter) so the file renders identically in Word on
any OS and in Google Docs — Charter only exists on macOS and silently falls back
elsewhere. Georgia is the closest universally-available transitional serif.

The YAML stays the source of record — backport docx edits that matter (CLAUDE.md).

Usage: uv run --with python-docx --with pyyaml scripts/yaml_to_docx.py <cv.yaml> <out.docx>
"""
import sys
import yaml
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn

FONT = "Georgia"
BODY = Pt(10)
CONTENT_WIDTH_IN = 7.5  # letter minus 0.5in margins

# RenderCV's month abbreviations (note June/July/Sept are spelled out)
MONTHS = ["", "Jan", "Feb", "Mar", "Apr", "May", "June", "July", "Aug", "Sept", "Oct", "Nov", "Dec"]


def fmt_date(d):
    if d is None:
        return None
    s = str(d)
    if s == "present":
        return "present"
    parts = s.split("-")
    if len(parts) >= 2:
        return f"{MONTHS[int(parts[1])]} {parts[0]}"
    return parts[0]


def date_range(e):
    if e.get("date"):
        return fmt_date(e["date"])
    start, end = fmt_date(e.get("start_date")), fmt_date(e.get("end_date") or "present")
    return f"{start} – {end}" if start else ""


def set_font(run, size=BODY, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), FONT)


def para(doc, before=0.0, after=0.0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    return p


def hrule(p, position="bottom", size="6"):
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.makeelement(qn("w:pBdr"), {})
    edge = pPr.makeelement(qn(f"w:{position}"), {})
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)      # eighths of a point
    edge.set(qn("w:space"), "3")
    edge.set(qn("w:color"), "000000")
    pbdr.append(edge)
    pPr.append(pbdr)


def section_title(doc, text):
    # harvard: centered bold title over a thin full-width rule
    p = para(doc, before=10, after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run(text), size=Pt(13), bold=True)
    hrule(p, "bottom", "4")


def entry_heading(doc, left_bold, left_rest, right):
    p = para(doc, before=8.5, after=2)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(CONTENT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT)
    set_font(p.add_run(left_bold), bold=True)
    if left_rest:
        set_font(p.add_run(left_rest))
    if right:
        set_font(p.add_run("\t" + right))


def bullet(doc, text):
    p = para(doc, after=1.7)  # ≈0.06cm — matches design.yaml rhythm
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.13)  # hanging bullet
    set_font(p.add_run("•  "))
    set_font(p.add_run(text))


def main(src, dst):
    cv = yaml.safe_load(open(src))["cv"]
    doc = Document()

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.5)
        section.left_margin = section.right_margin = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY

    # Header — name 25pt regular (harvard: not bold), contacts 9pt with • separators
    p = para(doc, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run(cv["name"]), size=Pt(25))

    contacts = [cv.get("location"), cv.get("email")]
    for s in cv.get("social_networks") or []:
        if s["network"] == "LinkedIn":
            contacts.append(f"linkedin.com/in/{s['username']}")
    p = para(doc, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_font(p.add_run("   •   ".join(c for c in contacts if c)), size=Pt(9))

    for key, entries in (cv.get("sections") or {}).items():
        section_title(doc, key.replace("_", " ").title() if key == key.lower() else key)

        for e in entries:
            if isinstance(e, str):
                p = para(doc, after=2)
                set_font(p.add_run(e))
            elif "company" in e:
                loc = f" – {e['location']}" if e.get("location") else ""
                entry_heading(doc, e["company"], f", {e['position']}{loc}", date_range(e))
                for h in e.get("highlights") or []:
                    bullet(doc, h)
            elif "institution" in e:
                degree = f"{e['degree']} in " if e.get("degree") else ""
                entry_heading(doc, e["institution"], f", {degree}{e['area']}", date_range(e))
                if e.get("summary"):
                    p = para(doc, after=1.7)
                    set_font(p.add_run(e["summary"]))
                for h in e.get("highlights") or []:
                    bullet(doc, h)
            elif "label" in e:
                p = para(doc, after=2.5)
                set_font(p.add_run(f"{e['label']}: "), bold=True)
                set_font(p.add_run(e["details"]))
            elif "bullet" in e:
                bullet(doc, e["bullet"])

    doc.save(dst)
    print(f"Wrote {dst}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        sys.exit(__doc__)
    main(sys.argv[1], sys.argv[2])
